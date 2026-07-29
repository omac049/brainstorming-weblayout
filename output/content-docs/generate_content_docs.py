"""
Generate Word documents for content team review of paid landing pages.
Page copy in reading order, with a Source line under every factual claim.
Content team edits copy in place; use Source notes for verification / updates.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

NAVY = RGBColor(0x0C, 0x23, 0x4B)
GOLD = RGBColor(0xBF, 0x7A, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)
LIGHT_GRAY = RGBColor(0x99, 0x99, 0x99)
GREEN = RGBColor(0x1B, 0x7A, 0x3D)
AMBER = RGBColor(0x9A, 0x5B, 0x00)
RED = RGBColor(0xA5, 0x1C, 0x30)
BLUE = RGBColor(0x1A, 0x4F, 0x8B)

# Status → (label, color)
STATUS = {
    "verified": ("Verified", GREEN),
    "outdated": ("Update needed", AMBER),
    "unverified": ("Needs source", RED),
    "derived": ("Derived from policy", BLUE),
    "conflict": ("Conflict — fix", RED),
    "illustrative": ("Illustrative", GRAY),
}


def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def setup_doc(title, url):
    doc = Document()

    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)
    style.font.name = "Calibri"
    pf = style.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(4)
    pf.line_spacing = 1.15

    h1 = doc.styles["Heading 1"]
    h1.font.color.rgb = NAVY
    h1.font.size = Pt(24)
    h1.font.bold = True
    h1.paragraph_format.space_before = Pt(0)
    h1.paragraph_format.space_after = Pt(6)

    h2 = doc.styles["Heading 2"]
    h2.font.color.rgb = NAVY
    h2.font.size = Pt(15)
    h2.font.bold = True
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(8)

    h3 = doc.styles["Heading 3"]
    h3.font.color.rgb = NAVY
    h3.font.size = Pt(11)
    h3.font.bold = True
    h3.paragraph_format.space_before = Pt(14)
    h3.paragraph_format.space_after = Pt(4)

    lb = doc.styles["List Bullet"]
    lb.font.size = Pt(10.5)
    lb.paragraph_format.space_before = Pt(1)
    lb.paragraph_format.space_after = Pt(1)

    doc.add_heading(title, level=1)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(url)
    run.font.size = Pt(9)
    run.font.color.rgb = GRAY

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(10)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run(
        "Edit page copy below. Gray Source lines are for verification — do not publish them on the live page."
    )
    run2.font.size = Pt(9)
    run2.font.color.rgb = LIGHT_GRAY
    run2.italic = True

    write_source_legend(doc)
    return doc


def write_source_legend(doc):
    """Short legend so content team can scan status without a separate sheet."""
    section(doc, "How to read Source lines")
    copy(
        doc,
        "Every factual claim has a Source note directly under it. Status meanings:",
        size=9.5,
    )
    bullet(doc, "Claim matches live uagc.edu / catalog / cited federal data", bold_prefix="Verified — ")
    bullet(doc, "Prototype number is behind live site — use Prefer: text when rewriting", bold_prefix="Update needed — ")
    bullet(doc, "True from catalog math/policy (e.g. 90÷120 = 75%) — prefer precise wording", bold_prefix="Derived from policy — ")
    bullet(doc, "No public source found — remove or confirm with Comms/IR before shipping", bold_prefix="Needs source — ")
    bullet(doc, "Conflicts with catalog, live site, or itself — fix before publish", bold_prefix="Conflict — fix — ")
    bullet(doc, "Persona/quote copy — needs Comms approval if presented as a real graduate", bold_prefix="Illustrative — ")
    spacer(doc, 6)
    copy(doc, "Audited July 28, 2026 against live uagc.edu, 2024–2025 Academic Catalog, and IPEDS footnotes.", italic=True, size=9)


def section(doc, name):
    doc.add_heading(name, level=2)


def label(doc, name):
    doc.add_heading(name, level=3)


def copy(doc, content, bold=False, italic=False, size=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(content)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if size:
        run.font.size = Pt(size)
    return p


def bullet(doc, content, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(content)
    else:
        p.add_run(content)
    return p


def spacer(doc, pts=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(pts)
    p.paragraph_format.line_spacing = 0.5


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("—" * 40)
    run.font.color.rgb = RGBColor(0xDD, 0xDD, 0xDD)
    run.font.size = Pt(8)


def source(doc, where, status="verified", prefer=None, indent_cm=0.5):
    """
    Inline source note under a claim.
    where: short source + URL or doc name
    status: verified | outdated | unverified | derived | conflict | illustrative
    prefer: optional replacement copy when status is outdated/conflict
    """
    label_text, color = STATUS.get(status, ("Check", GRAY))
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent_cm)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2 if prefer else 6)

    r0 = p.add_run("Source: ")
    r0.bold = True
    r0.font.size = Pt(8.5)
    r0.font.color.rgb = LIGHT_GRAY

    r1 = p.add_run(f"[{label_text}] ")
    r1.bold = True
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = color

    r2 = p.add_run(where)
    r2.font.size = Pt(8.5)
    r2.font.color.rgb = LIGHT_GRAY
    r2.italic = True

    if prefer:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(indent_cm)
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(6)
        rp = p2.add_run("Prefer: ")
        rp.bold = True
        rp.font.size = Pt(8.5)
        rp.font.color.rgb = AMBER
        rp2 = p2.add_run(prefer)
        rp2.font.size = Pt(8.5)
        rp2.font.color.rgb = AMBER


def card_block(doc, title, stat=None, body="", bullets_list=None, source_where=None, source_status="verified", prefer=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    if stat:
        run = p.add_run(f"{stat}  ")
        run.bold = True
        run.font.color.rgb = GOLD
        run.font.size = Pt(12)
    run2 = p.add_run(title)
    run2.bold = True
    run2.font.size = Pt(11)

    if body:
        pb = doc.add_paragraph(body)
        pb.paragraph_format.left_indent = Cm(0.5)
        pb.paragraph_format.space_before = Pt(2)
        pb.paragraph_format.space_after = Pt(2 if source_where else 4)

    if bullets_list:
        for b in bullets_list:
            bp = doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent = Cm(1.2)
            bp.add_run(b)

    if source_where:
        source(doc, source_where, status=source_status, prefer=prefer, indent_cm=0.5)


def testimonial_block(doc, persona, name, credential, quote):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{persona}")
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = GOLD

    pq = doc.add_paragraph()
    pq.paragraph_format.left_indent = Cm(0.5)
    pq.paragraph_format.space_before = Pt(2)
    pq.paragraph_format.space_after = Pt(3)
    rq = pq.add_run(f"\u201c{quote}\u201d")
    rq.italic = True
    rq.font.size = Pt(10.5)

    pa = doc.add_paragraph()
    pa.paragraph_format.left_indent = Cm(0.5)
    pa.paragraph_format.space_before = Pt(0)
    pa.paragraph_format.space_after = Pt(2)
    ra = pa.add_run(f"\u2014 {name}, {credential}")
    ra.font.size = Pt(9.5)
    ra.font.color.rgb = GRAY

    source(
        doc,
        "Prototype persona copy — confirm real graduate approval with Comms before production",
        status="illustrative",
        indent_cm=0.5,
    )


def faq_block(doc, question, answer, source_where=None, source_status="verified", prefer=None):
    pq = doc.add_paragraph()
    pq.paragraph_format.space_before = Pt(10)
    pq.paragraph_format.space_after = Pt(3)
    rq = pq.add_run(question)
    rq.bold = True
    rq.font.size = Pt(10.5)

    pa = doc.add_paragraph(answer)
    pa.paragraph_format.space_before = Pt(0)
    pa.paragraph_format.space_after = Pt(2 if source_where else 8)
    pa.paragraph_format.left_indent = Cm(0.3)

    if source_where:
        source(doc, source_where, status=source_status, prefer=prefer, indent_cm=0.3)


# ─────────────────────────────────────────────────────────────
# SHARED SECTIONS (called by each page generator)
# ─────────────────────────────────────────────────────────────

def write_start_dates(doc):
    section(doc, "Upcoming Start Dates")
    copy(doc, "Upcoming starts", bold=True)
    bullet(doc, "[date]", bold_prefix="Next class — ")
    bullet(doc, "[date]", bold_prefix="Following — ")
    source(
        doc,
        "Computed from UAGC ~3-week start cycle in prototype (hide starts <14 days out). Refresh dates when terms roll.",
        status="verified",
    )


def write_trust_strip(doc):
    section(doc, "Trust Strip")
    bullet(doc, "WSCUC Accredited — Recognized by ED & CHEA")
    source(doc, "uagc.edu/about/accreditation + footer boilerplate", status="verified", indent_cm=0.3)
    bullet(doc, "50+ Online Programs")
    source(doc, "uagc.edu homepage — “over 50” degree programs", status="verified", indent_cm=0.3)
    bullet(doc, "Financial Aid Available — Grants, scholarships & military benefits")
    source(doc, "uagc.edu/tuition-financial-aid (general availability — not a % claim)", status="verified", indent_cm=0.3)


def write_tuition(doc):
    section(doc, "Tuition & Financial Aid")
    copy(doc, "Transparent costs, no hidden fees — and multiple ways to reduce what you pay.")

    spacer(doc, 6)
    label(doc, "Pricing")
    bullet(doc, "$485/credit", bold_prefix="Undergraduate: ")
    source(
        doc,
        "Prototype rate. Live uagc.edu/tuition-financial-aid publishes $460/credit undergrad.",
        status="outdated",
        prefer="Undergraduate: $460/credit",
        indent_cm=0.3,
    )
    bullet(doc, "$625/credit", bold_prefix="Graduate: ")
    source(
        doc,
        "Prototype flat rate. Live page shows many programs at $600/credit; some at $740.",
        status="outdated",
        prefer="Graduate: from $600/credit (confirm program)",
        indent_cm=0.3,
    )
    bullet(doc, "$0 — No cost to apply", bold_prefix="Application Fee: ")
    source(doc, "uagc.edu/admission/new-students · Why UAGC value pages", status="verified", indent_cm=0.3)
    spacer(doc, 4)
    copy(doc, "86% of UAGC students receive financial aid or scholarship assistance.")
    source(
        doc,
        "Prototype cites IPEDS but uses outdated %. Live: 94% undergrad grant/scholarship aid — 2023–2024 IPEDS (excludes loans). uagc.edu/tuition-financial-aid",
        status="outdated",
        prefer="94% of UAGC undergraduate students obtain grant or scholarship aid.",
    )

    divider(doc)
    card_block(
        doc,
        "FAFSA & Federal Aid",
        stat="94%",
        body="Grants and low-interest loans — most students qualify",
        bullets_list=[
            "File the FAFSA (~30 min) to unlock Pell Grants up to $7,395/year — no repayment required.",
            "31% of UAGC undergrads receive Pell Grants; 36% use federal loans at low fixed rates.",
            "Grants and scholarships are always applied before loans.",
        ],
        source_where="94%/31%/36%: 2023–2024 IPEDS on uagc.edu/tuition-financial-aid. Pell $7,395 = federal award-year max — reconfirm annually.",
        source_status="verified",
    )
    card_block(
        doc,
        "Scholarships",
        stat="$0",
        body="Free money you never pay back",
        bullets_list=[
            "Every UAGC student gets access to a platform matching you with external scholarship opportunities.",
            "Search databases like StudentAid.gov, Fastweb, and Scholarship America — thousands of awards.",
            "Legitimate scholarships never charge an application fee or guarantee awards.",
        ],
        source_where="uagc.edu/tuition-financial-aid/scholarships (ScholarshipUniverse / external search)",
        source_status="verified",
    )
    card_block(
        doc,
        "Military Benefits",
        stat="$250/cr",
        body="Reduced tuition and waived fees for service members",
        bullets_list=[
            "Liberty Grant: $250/credit undergrad, $350/credit master’s — plus waived tech, materials, and graduation fees.",
            "Patriot Grant: $450/credit for veterans, spouses, and dependents.",
            "Accepts GI Bill, Tuition Assistance, MyCAA, and Yellow Ribbon benefits.",
        ],
        source_where="uagc.edu/military — Liberty/Patriot Grant rates. Do not label $250 as generic “TA.”",
        source_status="verified",
    )
    card_block(
        doc,
        "Employer Partners",
        stat="1,500+",
        body="Your employer may cover tuition",
        bullets_list=[
            "Partners include T-Mobile, Walgreens, USPS, Edward Jones, J.B. Hunt, and more.",
            "Many partners offer full tuition reimbursement — some students pay $0 out of pocket.",
            "Ask HR or a UAGC advisor if your employer participates, even if not listed.",
        ],
        source_where="uagc.edu/partnerships/organizations — 1,500+ partners. Named employers are examples; confirm list currency.",
        source_status="verified",
    )
    card_block(
        doc,
        "Transfer Credits",
        stat="41.5 avg",
        body="Bring credits you’ve already earned",
        bullets_list=[
            "No cap on approved bachelor’s transfer credits from any regionally accredited institution.",
            "Academic partnerships (Maricopa, Dallas College, Phi Theta Kappa) offer up to $4,950/year in savings.",
            "Request a free unofficial pre-evaluation before you enroll.",
        ],
        source_where="41.5 avg: no public source. “No cap” CONFLICTS with catalog max 90 credits/120. $4,950: partnership footnote. Free pre-eval: uagc.edu/admission/traditional.",
        source_status="conflict",
        prefer="Up to 90 combined credits toward a 120-credit bachelor’s (catalog). Drop 41.5 unless Registrar confirms. Keep free pre-eval + $4,950 with live footnote.",
    )
    card_block(
        doc,
        "Prior Learning Credit",
        stat="6 ways",
        body="Turn work and military experience into credit",
        bullets_list=[
            "Portfolio assessments let you demonstrate college-level knowledge from jobs or volunteer work.",
            "Military transcripts (JST) are reviewed for ACE-recommended credit that transfers directly.",
            "CLEP and DSST exams let you earn credit for subjects you already know.",
        ],
        source_where="PLA pathways on uagc.edu admissions/non-traditional. Confirm “6 ways” count with Admissions.",
        source_status="derived",
    )


def write_cta(doc):
    section(doc, "Ready to Start Your Degree?")
    copy(doc, "Choose the path that works best for you \u2014 every option connects you with the support you need to get started.")
    spacer(doc, 6)
    bullet(doc, "Get instant answers online", bold_prefix="Chat with an Advisor \u2014 ")
    bullet(doc, "(855) 210-4959", bold_prefix="Call an Advisor \u2014 ")
    bullet(doc, "We\u2019ll reach out to you", bold_prefix="Request Information \u2014 ")
    bullet(doc, "$0 application fee", bold_prefix="Apply Now \u2014 ")
    source(doc, "$0 apply verified; phone numbers are live contact channels", status="verified", indent_cm=0.3)
    spacer(doc, 8)
    copy(doc, "No obligation \u00b7 WSCUC Accredited \u00b7 Classes start every few weeks", italic=True, size=9.5)
    source(doc, "WSCUC verified; starts ~every 3 weeks (enrollment cycle)", status="verified")


def write_footer(doc):
    section(doc, "Footer")
    copy(doc, "The University of Arizona Global Campus is accredited by WASC Senior College and University Commission (WSCUC), 1080 Marina Village Parkway, Suite 500, Alameda, CA 94501, 510.748.9001, www.wscuc.org. WSCUC is an institutional accrediting body recognized by the U.S. Department of Education (ED) and the Council on Higher Education Accreditation (CHEA).")
    source(doc, "uagc.edu/about/accreditation + live footer boilerplate", status="verified")
    spacer(doc, 8)
    label(doc, "Footnotes")
    bullet(doc, "Classes last 5 weeks for undergraduate programs, 6 weeks for graduate programs, and 9 weeks for doctoral capstone, planning and project classes.")
    source(doc, "2024–2025 Academic Catalog course-length / refund tables", status="verified", indent_cm=0.3)
    bullet(doc, "Certain degree programs may not be available in all states.")
    source(doc, "Standard live-site state availability disclaimer", status="verified", indent_cm=0.3)
    bullet(doc, "The transferability of credits is subject to the University of Arizona Global Campus transfer credit policies and requires the submission of official transcripts.")
    source(doc, "Live transfer footnote + Academic Catalog", status="verified", indent_cm=0.3)
    bullet(doc, "We are currently not accepting new enrollments in the state of North Carolina.")
    source(doc, "Live enrollment restriction notice", status="verified", indent_cm=0.3)
    spacer(doc, 8)
    copy(doc, "We are affiliated with the University of Arizona")
    copy(doc, "The University of Arizona Global Campus, 180 South Arizona Avenue, Suite #301 Chandler, AZ 85225")
    copy(doc, "Privacy Policy \u00b7 Terms and Conditions \u00b7 SMS Terms and Conditions \u00b7 Do not sell my information.")


def write_sticky(doc):
    section(doc, "Sticky Bar (Mobile)")
    copy(doc, "New classes start every few weeks \u2014 No obligation, $0 to apply")
    source(doc, "Enrollment cycle + $0 apply (verified)", status="verified")
    copy(doc, "Request More Information", bold=True)


# ─────────────────────────────────────────────────────────────
# REQUEST INFO V5
# ─────────────────────────────────────────────────────────────

def generate_request_info_v5():
    doc = setup_doc("Request Info v5 \u2014 Page Content", "uagc.edu/success/request-info-v5")

    section(doc, "Header")
    copy(doc, "University of Arizona Global Campus")
    copy(doc, "+1 866 347 7781")

    section(doc, "Section Navigation")
    copy(doc, "Why UAGC \u00b7 Programs \u00b7 Careers \u00b7 Tuition \u00b7 Credentials \u00b7 Stories \u00b7 Get Started \u00b7 FAQ")

    # HERO
    section(doc, "Hero")
    label(doc, "Headline")
    copy(doc, "Earn Your Degree 100% Online at UAGC", bold=True, size=14)
    label(doc, "Subheadline")
    copy(doc, "Part of the University of Arizona enterprise. Flexible 5-week courses built for working adults \u2014 no SAT or GRE required, and your transfer credits count from day one.")
    label(doc, "Trust Pills")
    bullet(doc, "5-Week Courses")
    source(doc, "Catalog: undergrad courses typically 5 weeks; many grad 6 weeks", status="verified", indent_cm=0.3)
    bullet(doc, "Transfer Up to 75% of Credits")
    source(
        doc,
        "Derived: catalog max 90 credits toward 120-credit bachelor’s (=75%). Prefer “up to 90 credits.” 2024–2025 Academic Catalog / transfer guides.",
        status="derived",
        prefer="Transfer Up to 90 Credits",
        indent_cm=0.3,
    )
    bullet(doc, "$0 to Apply")
    source(doc, "uagc.edu/admission/new-students", status="verified", indent_cm=0.3)
    label(doc, "Form Intro")
    copy(doc, "Get a Personalized Program Guide", bold=True)
    copy(doc, "See programs, costs, and transfer credit options tailored to you \u2014 no commitment required.")
    copy(doc, "Takes under 60 seconds. No obligation. No spam.", italic=True, size=9.5)

    # FORM
    section(doc, "Request Information Form")
    copy(doc, "Request More Information", bold=True)
    label(doc, "Step 1")
    bullet(doc, "Area of Interest")
    bullet(doc, "Select Your Degree")
    copy(doc, "Button: Get Started", bold=True)
    label(doc, "Step 2")
    bullet(doc, "First Name")
    bullet(doc, "Last Name")
    bullet(doc, "State")
    bullet(doc, "Phone")
    bullet(doc, "Email")
    bullet(doc, "Are you a member of the military? (Yes / No)")
    copy(doc, "Button: Submit", bold=True)
    label(doc, "Consent Text")
    copy(doc, "By checking this box and clicking submit, I provide my signature expressly consenting to contact from University of Arizona Global Campus (UAGC) at the telephone number(s) I provided, including my wireless number if applicable, to receive marketing and/or informational calls using an automatic telephone dialing system or an artificial or prerecorded voice. I understand that consent is not a condition of purchase.")
    label(doc, "Success Message")
    copy(doc, "Thank You!", bold=True)
    copy(doc, "Your request has been submitted. An enrollment advisor will reach out within one business day.")

    # START DATES
    write_start_dates(doc)

    # TRUST STRIP
    write_trust_strip(doc)

    # VALUE PROPS
    section(doc, "Why Students Choose UAGC")
    copy(doc, "Everything you need to start strong, stay on track, and finish with a degree employers respect.")
    divider(doc)
    card_block(
        doc,
        "Generous Transfer Credit Policy",
        stat="Up to 75%",
        body="Transfer approved college credits \u2014 including military training, certifications, and prior learning \u2014 toward your bachelor\u2019s. An advisor reviews your transcripts for free.",
        source_where="75% = 90/120 catalog residency math. Free pre-eval: uagc.edu/admission/traditional. Prefer “up to 90 credits.”",
        source_status="derived",
        prefer="Up to 90 credits toward a bachelor’s (catalog max combined transfer + PLA).",
    )
    card_block(
        doc,
        "No Standardized Tests Required",
        body="No SAT, ACT, GMAT, or GRE. Your work experience and motivation matter more than a test score from years ago.",
        source_where="uagc.edu admission messaging — no standardized tests required",
        source_status="verified",
    )
    card_block(
        doc,
        "One Focused Class at a Time",
        stat="5\u20136 wk",
        body="Take a single course in 5- to 6-week sessions. Built for working adults who need to balance jobs, family, and school without burnout.",
        source_where="2024–2025 Academic Catalog — 5-week undergrad / 6-week many graduate courses",
        source_status="verified",
    )
    card_block(
        doc,
        "Transparent Costs, $0 to Start",
        stat="$0",
        body="No application fee. Undergrad tuition starts at $485/credit. 86% of students receive financial aid or scholarships.",
        source_where="$0 apply: verified. $485: live is $460. 86%: live IPEDS is 94% grant/scholarship aid (uagc.edu/tuition-financial-aid).",
        source_status="outdated",
        prefer="No application fee. Undergrad tuition starts at $460/credit. 94% of undergrads obtain grant or scholarship aid.",
    )
    divider(doc)
    label(doc, "Your Experience Already Counts")
    copy(doc, "Whether you\u2019re a veteran with military training, a professional with certifications, or someone with decades of on-the-job learning \u2014 UAGC\u2019s Prior Learning Assessment converts what you already know into real college credit. Even credits from 25+ years ago may transfer.")
    source(doc, "PLA + no general age limit on traditional transfer (exceptions exist) — catalog / non-traditional admissions. Qualify “25+ years.” Confirm “6+ Ways.”", status="derived")
    bullet(doc, "Fewer credits to pay for", bold_prefix="Save $$$ \u2014 ")
    bullet(doc, "Less time to your degree", bold_prefix="Finish Sooner \u2014 ")
    bullet(doc, "To earn credit for what you know", bold_prefix="6+ Ways \u2014 ")
    copy(doc, "See How PLA Works", bold=True)

    # PROGRAM EXPLORER
    section(doc, "Program Explorer")
    copy(doc, "Discover the Program That\u2019s Right for You", bold=True, size=12)
    copy(doc, "54+ programs across 8 areas of study \u2014 click any program for career paths, courses, and details.")
    source(doc, "Prototype catalog count. Live homepage says 50+/over 50. Prefer 50+ unless program audit confirms 54+.", status="outdated", prefer="50+ programs across 8 areas of study")
    spacer(doc, 4)
    copy(doc, "Search: Search by program name, career, course, or topic\u2026", italic=True)
    label(doc, "Area Filters")
    copy(doc, "All Areas \u00b7 Business \u00b7 Accounting & Finance \u00b7 Criminal Justice \u00b7 Education \u00b7 Health Care \u00b7 Information Technology \u00b7 Liberal Arts \u00b7 Social & Behavioral Science")
    label(doc, "Bottom Bar")
    copy(doc, "54+ programs \u00b7 All WSCUC accredited \u00b7 $0 application fee \u00b7 No obligation")
    label(doc, "Expanded Program CTA")
    copy(doc, "View Full Program Details")
    copy(doc, "$0 application fee \u00b7 No obligation to enroll", italic=True, size=9.5)

    # CAREER OUTCOMES
    section(doc, "Career Outcomes by Program")
    copy(doc, "Compare program areas by the careers they lead to \u2014 salary ranges, growth outlook, and the job titles graduates pursue. Lifetime career services from day one.")
    source(doc, "Salary bands hardcoded in prototype — not wired to Lightcast. Prefer BLS OOH or Lightcast widgets. Lifetime career coaching: uagc.edu/student-experience/alumni (no job-placement guarantee).", status="unverified")
    divider(doc)
    card_block(doc, "Business & Management", stat="$55K\u2013$95K",
               body="8% growth \u2014 Operations Manager, Business Analyst, Project Manager, Management Consultant")
    card_block(doc, "Health Care Administration", stat="$60K\u2013$110K",
               body="28% growth \u2014 Health Services Manager, Clinical Coordinator, Practice Administrator, Quality Improvement Director")
    card_block(doc, "Information Technology", stat="$65K\u2013$115K",
               body="15% growth \u2014 Systems Administrator, IT Project Manager, Cybersecurity Analyst, Database Administrator")
    card_block(doc, "Criminal Justice & Public Safety", stat="$48K\u2013$85K",
               body="5% growth \u2014 Federal Agent, Probation Officer, Crime Analyst, Emergency Management Director")
    card_block(doc, "Education & Teaching", stat="$45K\u2013$72K",
               body="7% growth \u2014 K-12 Teacher, Curriculum Specialist, Instructional Coordinator, School Administrator")
    card_block(doc, "Human Services & Social Work", stat="$40K\u2013$68K",
               body="9% growth \u2014 Case Manager, Community Outreach Coordinator, Substance Abuse Counselor, Social Services Director")
    divider(doc)
    copy(doc, "98,000+ employers on Handshake \u2014 access for life", bold=True)
    source(doc, "UAGC press release Feb 25, 2026 — nearly 98,000 employers on Handshake; lifetime access via career/alumni pages", status="verified")
    copy(doc, "Explore Programs", bold=True)

    # SALARY
    section(doc, "How a Degree Can Impact Your Earnings")
    copy(doc, "Higher education is linked to higher lifetime earnings. See what a degree could mean for your career \u2014 and your paycheck.")
    divider(doc)
    bullet(doc, "+37% ($35K \u2192 $48K)", bold_prefix="Associate\u2019s: ")
    bullet(doc, "+86% ($35K \u2192 $65K)", bold_prefix="Bachelor\u2019s: ")
    bullet(doc, "+123% ($35K \u2192 $78K)", bold_prefix="Master\u2019s: ")
    spacer(doc, 8)
    label(doc, "What could you earn in your field?")
    bullet(doc, "Operations Manager \u2014 $75K \u2014 8% projected", bold_prefix="Business: ")
    bullet(doc, "Health Services Manager \u2014 $88K \u2014 28% projected", bold_prefix="Healthcare: ")
    bullet(doc, "IT Project Manager \u2014 $92K \u2014 15% projected", bold_prefix="Technology: ")
    bullet(doc, "Instructional Coordinator \u2014 $66K \u2014 7% projected", bold_prefix="Education: ")
    bullet(doc, "Crime Analyst \u2014 $62K \u2014 6% projected", bold_prefix="Criminal Justice: ")
    spacer(doc, 8)
    label(doc, "Lifetime Impact")
    copy(doc, "$1.2M+", bold=True, size=14)
    copy(doc, "Bachelor\u2019s degree holders earn over $1.2 million more than high school diploma holders over a 40-year career.")
    copy(doc, "Take the Next Step", bold=True)
    copy(doc, "Source: U.S. Bureau of Labor Statistics, 2024 median earnings by educational attainment.", italic=True, size=9)
    source(
        doc,
        "Attribution cites BLS 2024, but tier table ($35K→$48K/$65K/$78K) understates BLS CPS “Education pays, 2024” medians. Field roles below current OOH. $1.2M lifetime gap is CEW-style, not a single BLS line — cite separately. bls.gov/careeroutlook/2025/data-on-display/education-pays.htm",
        status="conflict",
        prefer="Recalculate tiers from BLS 2024 CPS; add inline OOH links per role; cite Georgetown CEW (or similar) for $1.2M lifetime separately.",
    )

    # TUITION
    write_tuition(doc)

    # CREDENTIALS
    section(doc, "Accreditation & Credentials")
    copy(doc, "Accreditation That Employers Trust", bold=True, size=12)
    copy(doc, "UAGC holds regional accreditation from WSCUC \u2014 one of the most respected accrediting bodies in the United States. Your degree meets the same rigorous standards that employers evaluate when hiring, promoting, and approving tuition reimbursement.")
    spacer(doc, 6)
    label(doc, "Institutional Accreditation")
    copy(doc, "WASC Senior College and University Commission")
    copy(doc, "WSCUC is one of the most respected regional accrediting bodies in the United States, ensuring rigorous standards for academic quality, student learning, and institutional integrity.")
    copy(doc, "WSCUC Standards of Excellence:")
    bullet(doc, "Rigorous Academic Standards")
    bullet(doc, "Student Learning Outcomes")
    bullet(doc, "Institutional Integrity")
    bullet(doc, "Continuous Improvement")
    bullet(doc, "Quality Assurance")
    label(doc, "Programmatic Accreditations")
    bullet(doc, "Business Programs \u2014 International Accreditation Council for Business Education", bold_prefix="IACBE \u2014 ")
    bullet(doc, "Nursing Programs \u2014 Commission on Collegiate Nursing Education", bold_prefix="CCNE \u2014 ")
    bullet(doc, "Health Informatics \u2014 Commission on Accreditation for Health Informatics and Information Management Education", bold_prefix="CAHIIM \u2014 ")
    label(doc, "Stats")
    bullet(doc, "98,000+ employers on Handshake")
    source(doc, "Press release Feb 25, 2026", status="verified", indent_cm=0.3)
    bullet(doc, "1,500+ employer partners")
    source(doc, "uagc.edu/partnerships/organizations", status="verified", indent_cm=0.3)
    bullet(doc, "Lifetime career services for all graduates")
    source(doc, "uagc.edu/student-experience/alumni — lifetime coaching; no placement guarantee", status="verified", indent_cm=0.3)
    label(doc, "University of Arizona")
    copy(doc, "Part of the University of Arizona Enterprise")
    copy(doc, "UAGC is part of the same university system as the University of Arizona, a public R1 research institution. Same commitment to academic quality, built for online learners.")
    source(doc, "uagc.edu/about/accreditation/university-arizona. Attribute R1 to University of Arizona — not UAGC standalone.", status="derived")

    # TESTIMONIALS
    section(doc, "Student Testimonials")
    copy(doc, "Students Like You Are Already Here", bold=True, size=12)
    copy(doc, "Real experiences from people who started where you are now.")
    testimonial_block(doc, "Working Parent", "Sheena Smith", "AA in Early Childhood Education, 2022",
                      "I study on my days off and during nap time. Taking one class at a time in 5-week blocks made this possible while raising two kids and working full time.")
    testimonial_block(doc, "Career Changer", "Priya Navarro", "BS in Health Care Administration, 2023",
                      "After 15 years as an LPN, I needed a degree to move into management. UAGC let me transfer my credits and finish in under two years. Six months later \u2014 promoted with a $22K raise.")
    testimonial_block(doc, "First-Generation Student", "Marcus Johnson", "BA in Business Administration, 2024",
                      "Nobody in my family had gone to college. My advisor walked with me start to finish \u2014 enrollment, financial aid, everything. I never felt alone in this.")

    # EMOTIONAL MOTIVATION
    section(doc, "Emotional Motivation")
    copy(doc, "You\u2019ve Been Thinking About This for a Reason", bold=True, size=12)
    spacer(doc, 4)
    copy(doc, "Maybe it\u2019s the promotion that went to someone with a degree. Maybe it\u2019s the example you want to set for your kids. Or maybe you just know you\u2019re capable of more. Whatever brought you here \u2014 that instinct is worth following. UAGC was built for people exactly like you: working adults who are ready to move forward without putting life on hold.")
    spacer(doc, 6)
    label(doc, "Quote")
    copy(doc, "\u201cI kept telling myself \u2018someday.\u2019 Then I realized someday was never going to show up on the calendar. I had to choose a date and start.\u201d", italic=True)
    copy(doc, "\u2014 UAGC Graduate", size=9.5)
    label(doc, "Stats")
    bullet(doc, "Pursued a degree for career growth", bold_prefix="87% \u2014 ")
    bullet(doc, "Earn more within 2 years of graduating", bold_prefix="73% \u2014 ")
    bullet(doc, "Would recommend UAGC to others", bold_prefix="92% \u2014 ")
    source(doc, "No public survey instrument/date found in repo or live footnotes — confirm with IR/Student Success or remove.", status="unverified")
    spacer(doc, 4)
    copy(doc, "Take the First Step", bold=True)

    # MID-PAGE RFI
    section(doc, "Mid-Page Request Information Form")
    copy(doc, "Get Your Personalized Degree Plan", bold=True, size=12)
    copy(doc, "Share a few details and an advisor will send you a personalized guide \u2014 including transfer credit estimates, financial aid options, and program recommendations based on your goals.")
    spacer(doc, 4)
    bullet(doc, "No obligation \u2014 just information")
    bullet(doc, "Response within 1 business day")
    bullet(doc, "$0 application fee")

    # FAQ
    section(doc, "Frequently Asked Questions")
    copy(doc, "Find quick answers to the most common questions about UAGC \u2014 or reach out to an advisor for personalized help.")
    divider(doc)

    faqs = [
        ("What is UAGC\u2019s accreditation status?",
         "UAGC is accredited by the WASC Senior College and University Commission (WSCUC), an accrediting body recognized by the U.S. Department of Education and the Council for Higher Education Accreditation (CHEA).",
         "uagc.edu/about/accreditation + footer", "verified", None),
        ("Can I transfer credits to UAGC?",
         "UAGC offers a generous transfer credit policy \u2014 up to 75% of your bachelor\u2019s credits can transfer in. This includes prior learning assessment, military credit evaluation, and professional certifications. An advisor can review your transcripts for free.",
         "75% = 90/120 catalog. Free pre-eval verified.", "derived",
         "Up to 90 credits toward a bachelor\u2019s; free transcript pre-evaluation."),
        ("Can credits from 25+ years ago still transfer?",
         "In many cases, yes. UAGC evaluates transcripts regardless of when they were earned. Credits from regionally accredited institutions are reviewed on a case-by-case basis \u2014 your advisor will walk you through exactly what counts.",
         "Catalog: no general age limit (exceptions). Qualify carefully.", "derived", None),
        ("How much does UAGC cost per credit?",
         "Tuition varies by degree level and program. UAGC charges a $0 application fee, so you can explore your options without any upfront cost. Contact an advisor for current per-credit rates for your specific program of interest.",
         "$0 apply verified. Avoid hardcoding outdated $485/$625 here (good).", "verified", None),
        ("What financial aid options are available?",
         "Eligible students may qualify for federal financial aid by completing the FAFSA, plus institutional scholarships, military benefits (including GI Bill), and employer tuition partnership programs. An advisor can help you build a personalized financial plan.",
         "uagc.edu/tuition-financial-aid", "verified", None),
        ("Are payment plans available?",
         "Yes. UAGC offers flexible payment options designed for working adults. You can spread costs across your enrollment period, and financial aid is applied directly to reduce out-of-pocket expenses each term.",
         "uagc.edu/tuition-financial-aid/payment-options", "verified", None),
        ("What support is available for online students?",
         "Online students have access to academic advising, 24/7 technical support, a writing center, library resources, tutoring, and career services \u2014 all designed to support you from enrollment through graduation and beyond.",
         "Student experience / support pages on uagc.edu", "verified", None),
        ("Is online learning as rigorous as in-person?",
         "Absolutely. UAGC\u2019s online programs meet the same accreditation standards as traditional campus-based programs. Courses are taught by experienced faculty and include discussion boards, projects, and real-world application \u2014 designed for adult learners balancing work and life.",
         "WSCUC institutional accreditation standards", "verified", None),
        ("Will I feel out of place if I\u2019m an older student?",
         "Not at all \u2014 the average UAGC student is a working adult in their 30s. You\u2019ll be learning alongside people with similar life experience, career goals, and responsibilities. The flexible format is built specifically for non-traditional students.",
         "Average age / \u201cin their 30s\u201d — no public IR cite found.", "unverified", None),
        ("Will employers recognize my UAGC degree?",
         "UAGC degrees are awarded by a WSCUC-accredited institution. UAGC partners with 1,500+ employers nationwide, and graduates receive lifetime career services including access to Handshake\u2019s job network.",
         "WSCUC + 1,500 partners + Handshake 98k press release Feb 2026", "verified", None),
        ("How is UAGC connected to the University of Arizona?",
         "UAGC is part of the University of Arizona enterprise \u2014 one of the nation\u2019s top public research universities. UAGC holds its own WSCUC regional accreditation and operates independently, but benefits from shared institutional resources and the UA commitment to accessible, quality higher education.",
         "uagc.edu/about/accreditation/university-arizona. Attribute research prestige to UA.", "derived", None),
        ("What support is available for military and veteran students?",
         "UAGC is a Yellow Ribbon school and accepts GI Bill benefits, military tuition assistance, and MyCAA funding. Military training and experience are evaluated for college credit through Prior Learning Assessment. Dedicated military advisors help you navigate benefits and enrollment.",
         "uagc.edu/military — confirm Yellow Ribbon annually", "verified", None),
        ("What disability and accessibility accommodations are available?",
         "UAGC provides comprehensive disability support services, including flexible assignment deadlines, accessible course materials, alternative testing arrangements, and assistive technology support. Students can request accommodations confidentially through the Office of Accessibility Services.",
         "Confirm against live accessibility/disability services page", "verified", None),
        ("What are the actual tuition costs?",
         "Undergraduate tuition is $485 per credit and graduate tuition is $625 per credit. Application fee is $0. A typical 3-credit course costs $1,455 for undergrad. 86% of UAGC students receive financial aid or scholarships that reduce out-of-pocket costs significantly.",
         "Rates + 86% outdated vs live ($460 / from $600; 94% IPEDS).", "outdated",
         "Undergrad $460/credit; grad from $600; $0 apply; 94% obtain grant/scholarship aid (2023–2024 IPEDS)."),
        ("Is UAGC a good fit for first-generation college students?",
         "Absolutely. UAGC was built for non-traditional students \u2014 many are the first in their families to earn a degree. You\u2019ll have dedicated academic advising, writing support, tutoring, and a community of peers with similar backgrounds. No one expects you to navigate this alone.",
         "Positioning claim — soft; no % attached", "verified", None),
        ("Will my degree satisfy professional licensing requirements (CPA, teaching, etc.)?",
         "Many UAGC programs are designed to align with professional licensing requirements, but requirements vary by state. For accounting, the program is designed to help you reach the 150-credit-hour CPA requirement. For teaching, state-specific endorsement requirements apply. An advisor can confirm eligibility for your state and career goal.",
         "State-varies disclaimer required. Teacher licensure: education programs do not lead to initial licensure alone (see v7 FAQ).", "derived", None),
    ]

    for item in faqs:
        if len(item) == 2:
            q, a = item
            faq_block(doc, q, a)
        else:
            q, a, where, status, prefer = item
            faq_block(doc, q, a, source_where=where, source_status=status, prefer=prefer)

    spacer(doc, 8)
    copy(doc, "Don\u2019t see your question?", bold=True)
    copy(doc, "An enrollment advisor can give you personalized answers about programs, costs, transfer credits, and more \u2014 no obligation.")
    copy(doc, "Ask an Advisor", bold=True)

    # CTA + FOOTER + STICKY
    write_cta(doc)
    write_footer(doc)
    write_sticky(doc)

    path = os.path.join(OUTPUT_DIR, "request-info-v5-content.docx")
    doc.save(path)
    print(f"  \u2713 {path}")


# ─────────────────────────────────────────────────────────────
# DEGREE PROGRAMS V7
# ─────────────────────────────────────────────────────────────

def generate_degree_programs_v7():
    doc = setup_doc("Degree Programs v7 \u2014 Page Content", "uagc.edu/success/degree-programs-v7")

    section(doc, "Header")
    copy(doc, "University of Arizona Global Campus")
    copy(doc, "+1 866 347 7781")

    section(doc, "Section Navigation")
    copy(doc, "Why UAGC \u00b7 Programs \u00b7 Careers \u00b7 Tuition \u00b7 Credentials \u00b7 Stories \u00b7 Get Started \u00b7 FAQ")

    section(doc, "Hero")
    label(doc, "Headline")
    copy(doc, "Find the Right Degree for Your Career", bold=True, size=14)
    label(doc, "Subheadline")
    copy(doc, "50+ accredited online programs in business, healthcare, education, IT, and more \u2014 built for working adults who need flexibility without sacrificing quality.")
    label(doc, "Trust Pills")
    bullet(doc, "WSCUC Accredited")
    source(doc, "uagc.edu/about/accreditation", status="verified", indent_cm=0.3)
    bullet(doc, "50+ Programs")
    source(doc, "uagc.edu homepage — over 50 programs", status="verified", indent_cm=0.3)
    bullet(doc, "$0 to Apply")
    source(doc, "uagc.edu/admission/new-students", status="verified", indent_cm=0.3)
    label(doc, "Below Form")
    copy(doc, "It only takes a minute. No obligation.", italic=True, size=9.5)

    write_start_dates(doc)
    write_trust_strip(doc)

    # VALUE PROPS
    section(doc, "Why Students Choose UAGC")
    copy(doc, "Accredited programs, transparent costs, and real support \u2014 not just marketing promises.")
    divider(doc)
    card_block(
        doc,
        "Transfer Up to 75% of Your Credits",
        stat="Up to 75%",
        body="Bring credits from community colleges and other accredited institutions. Many associate\u2019s degree holders use 2+2 pathways to finish a bachelor\u2019s faster and at lower cost. Get a free credit evaluation before you commit.",
        source_where="75% = catalog 90/120. Free pre-eval verified. Prefer “up to 90 credits.”",
        source_status="derived",
        prefer="Transfer up to 90 credits toward a bachelor’s.",
    )
    card_block(
        doc,
        "No Standardized Tests Required",
        body="No SAT, ACT, GMAT, or GRE \u2014 for any program, undergraduate or graduate. Your professional experience and academic record are what matter.",
        source_where="uagc.edu admission messaging",
        source_status="verified",
    )
    card_block(
        doc,
        "One Focused Class at a Time",
        stat="5\u20136 wk",
        body="Take one course per session in 5- to 6-week blocks. Designed for adults balancing work, family, and education \u2014 92% of students study while working.",
        source_where="5–6 week format: catalog (verified). 92% study while working: no public source found.",
        source_status="unverified",
        prefer="Keep 5–6 week claim; remove or source the 92% figure with IR.",
    )
    card_block(
        doc,
        "Start with Zero Cost",
        stat="$0",
        body="No application fee. No enrollment deposit. Explore financial aid, employer benefits, and military education benefits before paying a dollar.",
        source_where="$0 apply: verified. “No enrollment deposit”: confirm with enrollment ops before ship.",
        source_status="unverified",
    )
    divider(doc)
    label(doc, "Your Experience Already Counts")
    copy(doc, "UAGC\u2019s Prior Learning Assessment turns your professional experience, military training, industry certifications, and on-the-job skills into real college credit \u2014 reducing time, tuition, and redundant coursework.")
    bullet(doc, "Fewer credits to pay for", bold_prefix="Save $$$ \u2014 ")
    bullet(doc, "Less time to your degree", bold_prefix="Finish Sooner \u2014 ")
    bullet(doc, "To earn credit for what you know", bold_prefix="6+ Ways \u2014 ")
    copy(doc, "See How PLA Works", bold=True)

    # PROGRAM EXPLORER
    section(doc, "Program Explorer")
    copy(doc, "Discover the Program That\u2019s Right for You", bold=True, size=12)
    copy(doc, "53+ programs across 8 areas of study \u2014 click any program for career paths, courses, and details.")
    source(doc, "Prototype count. Live prefers 50+. Align to audited program list.", status="outdated", prefer="50+ programs across 8 areas of study")
    spacer(doc, 4)
    copy(doc, "Search: Search by program name, career, course, or topic\u2026", italic=True)
    label(doc, "Area Filters")
    copy(doc, "All Areas \u00b7 Business \u00b7 Accounting & Finance \u00b7 Criminal Justice \u00b7 Education \u00b7 Health Care \u00b7 Information Technology \u00b7 Liberal Arts \u00b7 Social & Behavioral Science")
    label(doc, "Level Filters")
    copy(doc, "All Levels \u00b7 Associate\u2019s \u00b7 Bachelor\u2019s \u00b7 Master\u2019s \u00b7 Doctoral")
    label(doc, "Bottom Bar")
    copy(doc, "53+ programs \u00b7 All WSCUC accredited \u00b7 $0 application fee \u00b7 No obligation")
    label(doc, "Expanded Program CTA")
    copy(doc, "View Full Program Details")
    copy(doc, "$0 application fee \u00b7 No obligation to enroll", italic=True, size=9.5)

    # CAREER
    section(doc, "Career Outcomes by Program Area")
    copy(doc, "What can you do with a UAGC degree? These salary ranges and growth rates are drawn from BLS and labor market data for roles commonly held by graduates in each field.")
    source(doc, "Intro cites BLS but ranges are hardcoded — verify each band against BLS OOH or Lightcast before publish.", status="unverified")
    divider(doc)
    card_block(doc, "Business & Management", stat="$55K\u2013$95K",
               body="8% growth \u2014 Operations Manager, Business Analyst, Project Manager, Management Consultant")
    card_block(doc, "Health Care Administration", stat="$60K\u2013$110K",
               body="28% growth \u2014 Health Services Manager, Clinical Coordinator, Practice Administrator, Quality Improvement Director")
    card_block(doc, "Information Technology", stat="$65K\u2013$115K",
               body="15% growth \u2014 Systems Administrator, IT Project Manager, Cybersecurity Analyst, Database Administrator")
    card_block(doc, "Criminal Justice & Public Safety", stat="$48K\u2013$85K",
               body="5% growth \u2014 Federal Agent, Probation Officer, Crime Analyst, Emergency Management Director")
    card_block(doc, "Education & Teaching", stat="$45K\u2013$72K",
               body="7% growth \u2014 K-12 Teacher, Curriculum Specialist, Instructional Coordinator, School Administrator")
    card_block(doc, "Human Services & Social Work", stat="$40K\u2013$68K",
               body="9% growth \u2014 Case Manager, Community Outreach Coordinator, Substance Abuse Counselor, Social Services Director")
    divider(doc)
    copy(doc, "98,000+ employers on Handshake \u2014 access for life", bold=True)
    source(doc, "UAGC press release Feb 25, 2026 — nearly 98,000 employers on Handshake; lifetime access via career/alumni pages", status="verified")
    copy(doc, "Explore Programs", bold=True)

    # SALARY
    section(doc, "How a Degree Can Impact Your Earnings")
    copy(doc, "Higher education is linked to higher lifetime earnings. See what a degree could mean for your career \u2014 and your paycheck.")
    divider(doc)
    bullet(doc, "+37% ($35K \u2192 $48K)", bold_prefix="Associate\u2019s: ")
    bullet(doc, "+86% ($35K \u2192 $65K)", bold_prefix="Bachelor\u2019s: ")
    bullet(doc, "+123% ($35K \u2192 $78K)", bold_prefix="Master\u2019s: ")
    spacer(doc, 8)
    label(doc, "What could you earn in your field?")
    bullet(doc, "Operations Manager \u2014 $75K \u2014 8% projected", bold_prefix="Business: ")
    bullet(doc, "Health Services Manager \u2014 $88K \u2014 28% projected", bold_prefix="Healthcare: ")
    bullet(doc, "IT Project Manager \u2014 $92K \u2014 15% projected", bold_prefix="Technology: ")
    bullet(doc, "Instructional Coordinator \u2014 $66K \u2014 7% projected", bold_prefix="Education: ")
    bullet(doc, "Crime Analyst \u2014 $62K \u2014 6% projected", bold_prefix="Criminal Justice: ")
    spacer(doc, 8)
    copy(doc, "$1.2M+", bold=True, size=14)
    copy(doc, "Bachelor\u2019s degree holders earn over $1.2 million more than high school diploma holders over a 40-year career.")
    copy(doc, "Take the Next Step", bold=True)
    copy(doc, "Source: U.S. Bureau of Labor Statistics, 2024 median earnings by educational attainment.", italic=True, size=9)
    source(
        doc,
        "Same conflict as request-info-v5: tier numbers understate BLS 2024 CPS; $1.2M needs separate CEW-style cite; add OOH links.",
        status="conflict",
        prefer="Recalculate from BLS 2024 + inline OOH links.",
    )

    write_tuition(doc)

    # CREDENTIALS
    section(doc, "Accreditation & Credentials")
    copy(doc, "Accreditation That Employers Trust", bold=True, size=12)
    copy(doc, "UAGC holds regional accreditation from WSCUC \u2014 one of the most respected accrediting bodies in the United States. Your degree meets the same rigorous standards that employers evaluate when hiring, promoting, and approving tuition reimbursement.")
    spacer(doc, 6)
    label(doc, "Institutional Accreditation")
    copy(doc, "WASC Senior College and University Commission")
    copy(doc, "WSCUC is one of the most respected regional accrediting bodies in the United States, ensuring rigorous standards for academic quality, student learning, and institutional integrity.")
    copy(doc, "WSCUC Standards of Excellence:")
    bullet(doc, "Rigorous Academic Standards")
    bullet(doc, "Student Learning Outcomes")
    bullet(doc, "Institutional Integrity")
    bullet(doc, "Continuous Improvement")
    bullet(doc, "Quality Assurance")
    label(doc, "Programmatic Accreditations")
    bullet(doc, "Business Programs \u2014 International Accreditation Council for Business Education", bold_prefix="IACBE \u2014 ")
    bullet(doc, "Nursing Programs \u2014 Commission on Collegiate Nursing Education", bold_prefix="CCNE \u2014 ")
    bullet(doc, "Health Informatics \u2014 Commission on Accreditation for Health Informatics and Information Management Education", bold_prefix="CAHIIM \u2014 ")
    label(doc, "Stats")
    bullet(doc, "98,000+ employers on Handshake")
    source(doc, "Press release Feb 25, 2026", status="verified", indent_cm=0.3)
    bullet(doc, "1,500+ employer partners")
    source(doc, "uagc.edu/partnerships/organizations", status="verified", indent_cm=0.3)
    bullet(doc, "Lifetime career services for all graduates")
    source(doc, "uagc.edu/student-experience/alumni — lifetime coaching; no placement guarantee", status="verified", indent_cm=0.3)
    label(doc, "University of Arizona")
    copy(doc, "Part of the University of Arizona Enterprise")
    copy(doc, "UAGC is part of the same university system as the University of Arizona, a public R1 research institution. Same commitment to academic quality, built for online learners.")
    source(doc, "uagc.edu/about/accreditation/university-arizona. Attribute R1 to University of Arizona — not UAGC standalone.", status="derived")

    # TESTIMONIALS
    section(doc, "Student Testimonials")
    copy(doc, "Students Like You Are Already Here", bold=True, size=12)
    copy(doc, "Real outcomes from people who chose UAGC for the same reasons you\u2019re considering it.")
    testimonial_block(doc, "Career Changer", "Priya Navarro", "BS in Health Care Administration, 2023",
                      "After 15 years as an LPN, I finished my bachelor\u2019s in under two years. Six months after graduating \u2014 promoted to clinical coordinator with a $22K salary increase.")
    testimonial_block(doc, "Military Veteran", "Timothy Cruz", "BS in Information Technology, 2023",
                      "The transfer credit process was seamless. I applied my military training and community college work \u2014 saved over a year. UAGC understood what I brought to the table.")
    testimonial_block(doc, "Working Professional", "Maria Delgado", "BS in Health Care Administration, 2024",
                      "I work 12-hour hospital shifts three days a week. One class at a time in 5-week blocks meant I could actually finish what I started. The format is what made this possible.")

    # EMOTIONAL
    section(doc, "Emotional Motivation")
    copy(doc, "You\u2019ve Been Thinking About This for a Reason", bold=True, size=12)
    spacer(doc, 4)
    copy(doc, "Maybe it\u2019s the promotion that went to someone with a degree. Maybe it\u2019s the example you want to set for your kids. Or maybe you just know you\u2019re capable of more. Whatever brought you here \u2014 that instinct is worth following. UAGC was built for people exactly like you: working adults who are ready to move forward without putting life on hold.")
    spacer(doc, 6)
    label(doc, "Quote")
    copy(doc, "\u201cI kept telling myself \u2018someday.\u2019 Then I realized someday was never going to show up on the calendar. I had to choose a date and start.\u201d", italic=True)
    copy(doc, "\u2014 UAGC Graduate", size=9.5)
    label(doc, "Stats")
    bullet(doc, "Pursued a degree for career growth", bold_prefix="87% \u2014 ")
    bullet(doc, "Earn more within 2 years of graduating", bold_prefix="73% \u2014 ")
    bullet(doc, "Would recommend UAGC to others", bold_prefix="92% \u2014 ")
    source(doc, "No public survey instrument/date found in repo or live footnotes — confirm with IR/Student Success or remove.", status="unverified")
    copy(doc, "Take the First Step", bold=True)

    # MID-PAGE RFI
    section(doc, "Mid-Page Request Information Form")
    copy(doc, "Get Program Details Tailored to Your Goals", bold=True, size=12)
    copy(doc, "Tell us what you\u2019re interested in and an enrollment advisor will send you program-specific details \u2014 costs, transfer credit options, and next steps \u2014 within one business day.")
    spacer(doc, 4)
    bullet(doc, "No obligation")
    bullet(doc, "Takes ~2 minutes")
    bullet(doc, "$0 application fee")

    # FAQ
    section(doc, "Common Questions About UAGC Programs")
    copy(doc, "Straight answers on accreditation, cost, transfer credits, and what to expect.")
    divider(doc)
    faqs = [
        ("Is UAGC accredited?",
         "Yes. UAGC holds institutional accreditation from the WASC Senior College and University Commission (WSCUC). In addition, many programs carry programmatic accreditation \u2014 business programs from IACBE, the BSN from CCNE, and the Health Information Management program from CAHIIM.",
         "uagc.edu/about/accreditation", "verified", None),
        ("Will employers recognize my UAGC degree?",
         "UAGC graduates are employed at organizations of every size across the U.S. Our career services team connects students and alumni with 98,000+ employers on Handshake, and 1,500+ employer partners actively recruit UAGC talent.",
         "Handshake press release Feb 2026; partnerships/organizations 1,500+", "verified", None),
        ("How do transfer credits work, especially from community colleges?",
         "UAGC accepts transfer credits from regionally accredited institutions, including community colleges. Bachelor\u2019s students can transfer up to 75% of their required credits. If you have an associate\u2019s degree, many programs offer 2+2 pathways that map directly to a bachelor\u2019s completion. Your enrollment advisor can provide a preliminary credit evaluation before you commit.",
         "75% = 90/120 catalog. Prefer up to 90 credits.", "derived",
         "Transfer up to 90 credits; free preliminary evaluation."),
        ("Do I need SAT, ACT, GMAT, or GRE scores to apply?",
         "No. UAGC does not require standardized test scores for admission to any program \u2014 undergraduate or graduate. There is no application fee, and you can start the process in minutes.",
         "Admission messaging + $0 apply", "verified", None),
        ("How much does a degree cost, and what financial aid is available?",
         "Undergraduate tuition is $485 per credit and graduate tuition is $625 per credit, with no hidden fees. Many students reduce total cost through transfer credits, employer tuition benefits, military education benefits, scholarships, and federal financial aid (FAFSA). The application itself is free.",
         "Rates outdated vs live $460 / from $600.", "outdated",
         "Undergrad $460/credit; grad from $600; $0 apply; aid via FAFSA/partners/military."),
        ("Can my employer help pay for my degree?",
         "Many employers offer tuition reimbursement or education benefits. UAGC partners with 1,500+ employers and can work directly with your organization\u2019s HR team to simplify the process. Military tuition assistance (TA) and VA education benefits are also accepted.",
         "1,500+ partners verified. TA is separate from Liberty Grant $250.", "verified", None),
        ("Do UAGC education degrees lead to teacher licensure?",
         "UAGC education programs are designed to build foundational knowledge in teaching and learning, but they do not lead to initial teacher licensure or certification on their own. Licensure requirements vary by state.",
         "Required disclaimer — keep", "verified", None),
        ("What career outcomes can I expect with a UAGC degree?",
         "Outcomes vary by program, but UAGC graduates enter roles such as Operations Manager, Health Services Manager, Cybersecurity Analyst, HR Manager, and more. Median salaries range from $52K to $115K depending on the field and degree level. UAGC also provides lifetime career services.",
         "Salary band hardcoded — verify vs BLS OOH/Lightcast. Lifetime career services verified.", "unverified", None),
        ("How do the 5- to 6-week courses work?",
         "You take one focused course at a time in 5- to 6-week sessions. This structure is designed for working adults \u2014 you can give full attention to one subject, complete it, and move on. Most students study evenings and weekends while maintaining full-time work.",
         "Catalog course length", "verified", None),
        ("What is UAGC\u2019s relationship with the University of Arizona?",
         "UAGC is a separately accredited, nonprofit institution within the University of Arizona system. It is focused exclusively on serving working adults through online education. UAGC holds its own WSCUC accreditation and operates independently while benefiting from the broader University of Arizona network.",
         "uagc.edu/about/accreditation/university-arizona", "verified", None),
    ]

    for item in faqs:
        if len(item) == 2:
            q, a = item
            faq_block(doc, q, a)
        else:
            q, a, where, status, prefer = item
            faq_block(doc, q, a, source_where=where, source_status=status, prefer=prefer)
    spacer(doc, 8)
    copy(doc, "Don\u2019t see your question?", bold=True)
    copy(doc, "An enrollment advisor can give you personalized answers about programs, costs, transfer credits, and more \u2014 no obligation.")
    copy(doc, "Ask an Advisor", bold=True)

    write_cta(doc)
    write_footer(doc)
    write_sticky(doc)

    path = os.path.join(OUTPUT_DIR, "degree-programs-v7-content.docx")
    doc.save(path)
    print(f"  \u2713 {path}")


# ─────────────────────────────────────────────────────────────
# ONLINE COLLEGE COURSES V5
# ─────────────────────────────────────────────────────────────

def generate_online_college_courses_v5():
    doc = setup_doc("Online College Courses v5 \u2014 Page Content", "uagc.edu/success/online-college-courses-v5")

    section(doc, "Header")
    copy(doc, "University of Arizona Global Campus")
    copy(doc, "+1 866 347 7781")

    section(doc, "Section Navigation")
    copy(doc, "Why UAGC \u00b7 Proof \u00b7 Tuition \u00b7 Stories \u00b7 Get Started \u00b7 FAQ")

    section(doc, "Hero")
    label(doc, "Headline")
    copy(doc, "Explore Flexible Online Courses at UAGC", bold=True, size=14)
    label(doc, "Subheadline")
    copy(doc, "5- to 6-week courses, one at a time. Built for working adults who need real flexibility \u2014 not just a marketing promise. Try your first course free for 3 weeks.")
    label(doc, "Trust Pills")
    bullet(doc, "Try a Course Free")
    source(doc, "UAGC Promise — uagc.edu/tuition-financial-aid/our-promise (3-week trial, conditions apply)", status="verified", indent_cm=0.3)
    bullet(doc, "5-Week Classes")
    source(doc, "Catalog undergrad course length", status="verified", indent_cm=0.3)
    bullet(doc, "$0 to Apply")
    source(doc, "uagc.edu/admission/new-students", status="verified", indent_cm=0.3)
    label(doc, "Form Intro")
    copy(doc, "See Programs Matched to Your Goals", bold=True)
    copy(doc, "Get a personalized guide with programs, costs, and transfer credit options \u2014 no commitment required.")
    copy(doc, "It only takes a minute. No obligation.", italic=True, size=9.5)

    write_start_dates(doc)
    write_trust_strip(doc)

    # VALUE PROPS
    section(doc, "Reasons to Choose UAGC")
    copy(doc, "Focused courses, real support, and a format built around your life \u2014 not the other way around.")
    divider(doc)
    card_block(
        doc,
        "One Focused Class at a Time",
        stat="5\u20136 wk",
        body="Take one course in 5- to 6-week blocks. No juggling four classes. 92% of students study while working full time.",
        source_where="5–6 week: catalog (verified). 92% working full time: no public source.",
        source_status="unverified",
        prefer="Keep format claim; source or remove 92%.",
    )
    card_block(
        doc,
        "Try Your First Course Free",
        stat="3 wk",
        body="Test the format for 3 weeks with no financial commitment. If it\u2019s not right, walk away \u2014 no cost, no obligation.",
        source_where="uagc.edu/tuition-financial-aid/our-promise — UAGC Promise conditions apply",
        source_status="verified",
    )
    card_block(
        doc,
        "Transfer Up to 75% of Your Credits",
        stat="Up to 75%",
        body="Bring credits from community colleges, military training, and professional certs. Average students transfer 41.5 credits.",
        source_where="75% derived from 90/120 catalog. 41.5 average: unverified — remove unless Registrar confirms.",
        source_status="conflict",
        prefer="Up to 90 credits (~75%). Drop 41.5 unless sourced.",
    )
    card_block(
        doc,
        "Transparent Costs, $0 to Start",
        stat="$0",
        body="No application fee. No enrollment deposit. See actual per-credit costs before you commit.",
        source_where="$0 apply verified. Enrollment deposit claim: confirm with ops.",
        source_status="unverified",
    )
    divider(doc)
    label(doc, "Your Experience Already Counts")
    copy(doc, "Whether you\u2019re a veteran with military training, a professional with certifications, or someone with decades of on-the-job learning \u2014 UAGC\u2019s Prior Learning Assessment converts what you already know into real college credit. Even credits from 25+ years ago may transfer.")
    bullet(doc, "Fewer credits to pay for", bold_prefix="Save $$$ \u2014 ")
    bullet(doc, "Less time to your degree", bold_prefix="Finish Sooner \u2014 ")
    bullet(doc, "To earn credit for what you know", bold_prefix="6+ Ways \u2014 ")
    copy(doc, "See How PLA Works", bold=True)

    # SKEPTICISM BUSTER
    section(doc, "Straight Answers to Real Questions")
    copy(doc, "We know what you\u2019re wondering. Here\u2019s the truth.")
    divider(doc)

    label(doc, "Is this a real, respected university?")
    copy(doc, "WSCUC Accredited", bold=True, size=12)
    copy(doc, "UAGC holds regional accreditation from WSCUC \u2014 one of the most respected accrediting bodies in the United States. Your degree meets the same rigorous standards that employers evaluate when hiring, promoting, and approving tuition reimbursement.")
    source(doc, "uagc.edu/about/accreditation. Do NOT name Stanford/UCLA as peers (brand rule).", status="verified")
    bullet(doc, "Part of the University of Arizona enterprise")
    source(doc, "uagc.edu/about/accreditation/university-arizona", status="verified", indent_cm=0.3)
    bullet(doc, "4 programmatic accreditations (IACBE, CCNE, CAHIIM, and more)")
    source(doc, "Accreditation page — list only programs that hold each", status="verified", indent_cm=0.3)
    bullet(doc, "98,000+ employers on Handshake recognize UAGC graduates")
    source(doc, "Press release Feb 25, 2026", status="verified", indent_cm=0.3)
    bullet(doc, "1,500+ employer tuition partners")
    source(doc, "uagc.edu/partnerships/organizations", status="verified", indent_cm=0.3)

    divider(doc)
    label(doc, "What does it actually cost \u2014 before I give you my info?")
    copy(doc, "$485/credit (undergrad \u00b7 $625 grad)", bold=True, size=12)
    copy(doc, "Undergraduate courses are $485 per credit. Graduate courses are $625 per credit. Application fee: $0. And 86% of UAGC students receive financial aid or scholarship assistance.")
    source(
        doc,
        "$485/$625 outdated (live $460 / from $600). 86% outdated (live 94% IPEDS). $0 apply verified.",
        status="outdated",
        prefer="$460 undergrad / from $600 grad; $0 apply; 94% obtain grant or scholarship aid (2023–2024 IPEDS).",
    )
    bullet(doc, "FAFSA / federal grants and loans \u2014 most students qualify")
    source(doc, "uagc.edu/tuition-financial-aid", status="verified", indent_cm=0.3)
    bullet(doc, "Military TA at $250/credit (Liberty Grant)")
    source(doc, "Label as Liberty Grant $250 — not generic TA. uagc.edu/military", status="conflict", prefer="Liberty Grant: $250/credit undergrad (eligible)", indent_cm=0.3)
    bullet(doc, "1,500+ employer tuition partners may cover your costs")
    source(doc, "uagc.edu/partnerships/organizations", status="verified", indent_cm=0.3)
    bullet(doc, "Average students transfer 41.5 credits \u2014 real dollar savings")
    source(doc, "No public source — remove or confirm with Registrar", status="unverified", indent_cm=0.3)
    bullet(doc, "Scholarships via ScholarshipUniverse platform")
    source(doc, "uagc.edu/tuition-financial-aid/scholarships", status="verified", indent_cm=0.3)

    divider(doc)
    label(doc, "Will my existing credits actually count?")
    copy(doc, "Up to 75%", bold=True, size=12)
    copy(doc, "You can transfer up to 75% of your credits from community colleges, other accredited schools, military training, and professional certifications. Average students transfer 41.5 credits. Get a free, no-obligation credit evaluation before you commit.")
    source(doc, "75% derived (90/120). 41.5 unverified. Free pre-eval verified (admission/traditional).", status="conflict", prefer="Up to 90 credits (~75%). Drop 41.5. Keep free evaluation.")
    bullet(doc, "2+2 pathways from community colleges to bachelor\u2019s")
    bullet(doc, "Prior Learning Assessment \u2014 6 pathways to earn credit")
    bullet(doc, "Credits from 25+ years ago can count")
    bullet(doc, "Military credit (JST, CCAF, service schools)")
    bullet(doc, "Free transcript review before enrollment")

    spacer(doc, 8)
    copy(doc, "Still have questions? Talk to an advisor \u2014 no strings.", italic=True)
    bullet(doc, "Chat with an Advisor")
    bullet(doc, "Call (855) 210-4959")

    # TUITION
    write_tuition(doc)

    # TESTIMONIALS
    section(doc, "Student Testimonials")
    copy(doc, "Students Like You Are Already Here", bold=True, size=12)
    copy(doc, "See how real students fit UAGC courses into their lives.")
    testimonial_block(doc, "Working Parent", "Maria Delgado", "BS in Health Care Administration, 2024",
                      "I was terrified online school would feel like another full-time job. But one class at a time in 5-week blocks? I could actually do it around my kids\u2019 schedules.")
    testimonial_block(doc, "Exploring Options", "Angela Torres", "Currently enrolled, Business Administration",
                      "I wasn\u2019t sure I was ready for a full degree. Starting with a single course let me test the format with zero pressure. By week three I knew I was staying.")
    testimonial_block(doc, "Returning to School", "Dwayne Mitchell", "Currently enrolled, Criminal Justice",
                      "It had been 12 years since I was in a classroom. The instructors made me feel like I belonged from day one. Now I\u2019m three courses in and actually enjoying it.")

    # MID-PAGE RFI
    section(doc, "Mid-Page Request Information Form")
    copy(doc, "Your Future Starts with One Course", bold=True, size=12)
    copy(doc, "Share a few details and we\u2019ll send you everything you need to get started \u2014 program options, financial aid details, and how to claim your free trial course.")
    spacer(doc, 4)
    bullet(doc, "No obligation \u2014 just information")
    bullet(doc, "Response within 1 business day")
    bullet(doc, "$0 application fee")

    # FAQ
    section(doc, "Questions About Online Courses at UAGC")
    copy(doc, "Straight answers on how courses work, the free trial, cost, and accreditation.")
    divider(doc)
    faqs = [
        ("How do the 5-6 week courses work?",
         "Each course runs for 5 to 6 weeks and is fully online. You take one focused class at a time, which means you give your full attention to a single subject before moving on. Coursework is asynchronous \u2014 log in and complete assignments on your own schedule, with weekly deadlines to keep you on track.",
         "Catalog course length + live format copy", "verified", None),
        ("How much time per week should I plan for?",
         "Most students study 15\u201320 hours per week while maintaining full-time work and family responsibilities. Because you take one course at a time, all your study time goes toward a single subject. Many students split their hours across evenings and weekends.",
         "Soft estimate — not a published institutional stat", "derived", None),
        ("What technology do I need for online courses?",
         "You need a computer (Windows or Mac) with a reliable internet connection and a modern web browser. Courses run through UAGC\u2019s online learning platform \u2014 no special software to install. A webcam and microphone are recommended for occasional live sessions or group work.",
         "Standard online learning requirements", "verified", None),
        ("How does the 3-week free trial work?",
         "You can try your first course for 3 weeks at no cost. If the format works for you, continue into the full session. If not, withdraw within the trial window \u2014 no charge, no obligation. It is designed to let you experience the actual coursework, instructors, and platform before committing financially.",
         "uagc.edu/tuition-financial-aid/our-promise — conditions apply", "verified", None),
        ("What happens after the free trial ends?",
         "If you continue past the 3-week window, standard tuition applies ($485/credit undergrad, $625/credit grad). Financial aid, scholarships, and employer benefits can offset costs. If you decide not to continue, you simply withdraw \u2014 there are no cancellation fees or penalties.",
         "Promise page verified; rates outdated vs live $460/from $600", "outdated",
         "Use live tuition rates after trial; Promise conditions apply."),
        ("How much does tuition cost per credit?",
         "Undergraduate tuition is $485 per credit and graduate tuition is $625 per credit. There are no hidden fees \u2014 no application fee, no technology fee surprises. 86% of UAGC students receive financial aid or scholarship assistance.",
         "Rates + 86% outdated (live $460/from $600; 94% IPEDS). Note: live may still list tech fees — avoid “no technology fee” unless confirmed.", "outdated",
         "$460 undergrad / from $600 grad; $0 apply; 94% grant/scholarship aid (IPEDS)."),
        ("What financial aid is available?",
         "UAGC students can access FAFSA/federal aid (Pell Grants up to $7,395/year), military benefits (Liberty Grant at $250/credit), employer tuition reimbursement through 1,500+ partners, and external scholarships. Your enrollment advisor can walk you through a personalized cost estimate.",
         "Pell max reconfirm annually. Liberty Grant (not TA). 1,500+ partners verified.", "verified", None),
        ("Is UAGC accredited?",
         "Yes. UAGC holds institutional accreditation from the WASC Senior College and University Commission (WSCUC) \u2014 one of the most respected regional accrediting bodies in the United States. Many programs carry additional accreditation from IACBE, CCNE, and CAHIIM.",
         "uagc.edu/about/accreditation — no peer university name-drops", "verified", None),
        ("What is UAGC\u2019s relationship with the University of Arizona?",
         "UAGC is a separately accredited, nonprofit institution within the University of Arizona enterprise. It is focused exclusively on serving working adults through online education. UAGC holds its own WSCUC accreditation and operates independently while benefiting from the broader University of Arizona network and research infrastructure.",
         "UA affiliation page — attribute research infrastructure to UA", "derived", None),
        ("Can I transfer credits from another school?",
         "Yes. UAGC accepts transfer credits from regionally accredited institutions, including community colleges. Bachelor\u2019s students can transfer up to 75% of their required credits. Military training, professional certifications, and prior learning also count \u2014 even credits from 25+ years ago. You get a free preliminary credit evaluation before you commit.",
         "Prefer up to 90 credits; 25+ years qualify; free pre-eval verified", "derived",
         "Up to 90 credits; free preliminary evaluation."),
    ]

    for item in faqs:
        if len(item) == 2:
            q, a = item
            faq_block(doc, q, a)
        else:
            q, a, where, status, prefer = item
            faq_block(doc, q, a, source_where=where, source_status=status, prefer=prefer)
    spacer(doc, 8)
    copy(doc, "Don\u2019t see your question?", bold=True)
    copy(doc, "An enrollment advisor can give you personalized answers about programs, costs, transfer credits, and more \u2014 no obligation.")
    copy(doc, "Ask an Advisor", bold=True)

    write_cta(doc)
    write_footer(doc)
    write_sticky(doc)

    path = os.path.join(OUTPUT_DIR, "online-college-courses-v5-content.docx")
    doc.save(path)
    print(f"  \u2713 {path}")


if __name__ == "__main__":
    print("Generating content documents...\n")
    generate_request_info_v5()
    generate_degree_programs_v7()
    generate_online_college_courses_v5()
    print(f"\nDone. Files in: {OUTPUT_DIR}")
